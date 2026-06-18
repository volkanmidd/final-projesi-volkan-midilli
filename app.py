"""
Yerel Belge Soru-Cevap ve Model Karşılaştırma Asistanı (Seçenek 4)

Bu uygulama Ollama üzerinden yerel olarak çalışan dil modellerini (qwen2.5:1.5b, phi3 vb.)
kullanarak yüklenen dokümanlar (PDF, DOCX, TXT) üzerinden analiz yapar.
Veriler tamamen yerel cihazda işlendiği için tam gizlilik sağlar.
"""

import os
import io
import time
import hashlib
import numpy as np
import streamlit as st
from PIL import Image
import pytesseract
from typing import List

# LangChain bileşenleri
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.llms import Ollama
from langchain.prompts import PromptTemplate
from langchain_core.embeddings import Embeddings
from sklearn.feature_extraction.text import TfidfVectorizer

# Tesseract OCR Yolu (Kendi bilgisayarındaki yola göre güncelleyebilirsin)
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

# ----------------- SAYFA AYARLARI ----------------- #
st.set_page_config(
    page_title="Offline Kullanılabilen Yapay Zeka Doküman Asistanı",
    page_icon="🤖",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------------- SABİTLER ----------------- #
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200
SIMILARITY_TOP_K = 4

# ----------------- ÖZELLEŞTİRİLMİŞ STİL (CSS) ----------------- #
st.markdown(
    """
<style>
    .main-header {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
        box-shadow: 0 4px 15px rgba(0,0,0,0.1);
    }
    .feature-card {
        background: white;
        padding: 1.5rem;
        border-radius: 10px;
        box-shadow: 0 2px 10px rgba(0,0,0,0.05);
        border-left: 4px solid #1e3c72;
        margin: 1rem 0;
        color: #333;
    }
    .stButton > button {
        background: linear-gradient(135deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.5rem 2rem;
        font-weight: bold;
    }
</style>
""",
    unsafe_allow_html=True,
)


# ----------------- YEREL TF-IDF EMBEDDINGS ----------------- #
class TFIDFEmbeddings(Embeddings):
    """Tamamen yerel çalışan, internet bağımlılığı olmayan gömme (embedding) sınıfı."""

    def __init__(self, max_features: int = 384):
        self.vectorizer = TfidfVectorizer(
            max_features=max_features,
            stop_words="english",
            ngram_range=(1, 2),
            lowercase=True,
            token_pattern=r"\b[a-zA-Z]{2,}\b",
        )
        self.is_fitted = False
        self.dimension = max_features

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        if not self.is_fitted:
            self.vectorizer.fit(texts)
            self.is_fitted = True
        vectors = self.vectorizer.transform(texts).toarray()
        return [self._pad(v) for v in vectors]

    def embed_query(self, text: str) -> List[float]:
        if not self.is_fitted:
            return [0.0] * self.dimension
        vec = self.vectorizer.transform([text]).toarray()[0]
        return self._pad(vec)

    def _pad(self, vector: np.ndarray) -> List[float]:
        if len(vector) < self.dimension:
            vector = np.pad(vector, (0, self.dimension - len(vector)), "constant")
        return vector[: self.dimension].tolist()


# ----------------- DOKÜMAN YÖNETİCİSİ ----------------- #
class DocumentManager:
    def __init__(self):
        self.documents: List[Document] = []
        self.processed_files = {}
        self.embeddings = TFIDFEmbeddings()
        self.vectordb = None

        # Yönerge Madde 5'e uygun Prompt Tasarımı (Prompt Tuning)
        self.prompt_template = PromptTemplate(
            input_variables=["context", "question"],
            template="""Sen yüklenen dokümanlara göre cevap veren uzman bir yerel yapay zekâ asistanısın. 
Aşağıdaki bağlamı (context) kullanarak soruyu kapsamlı ve doğru bir şekilde cevapla. 
Eğer cevap bağlam içerisinde yoksa, kesinlikle bilgi uydurma ve tam olarak şu cümleyi söyle: "Verilen dokümanlarda bu sorunun cevabını bulamadım."

Bağlam:
{context}

Soru: {question}

Cevap:""",
        )

    def add_file(self, filename: str, content: str, file_hash: str, file_size: int):
        if file_hash in self.processed_files:
            return False, f"'{filename}' zaten yüklenmiş (Mükerrer dosya tespiti)."
        if not content.strip():
            return False, f"'{filename}' içeriği boş veya okunamadı."

        doc = Document(page_content=content,
                       metadata={"source": filename, "file_hash": file_hash, "file_size": file_size})
        self.documents.append(doc)
        self.processed_files[file_hash] = {"name": filename, "size": file_size, "word_count": len(content.split())}
        return True, f"✅ '{filename}' başarıyla işlendi ({len(content.split())} kelime)."

    def _rebuild_vectordb(self):
        if not self.documents:
            self.vectordb = None
            return
        try:
            splitter = RecursiveCharacterTextSplitter(chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP)
            chunks = splitter.split_documents(self.documents)
            all_texts = [c.page_content for c in chunks]
            self.embeddings = TFIDFEmbeddings()
            _ = self.embeddings.embed_documents(all_texts)
            self.vectordb = FAISS.from_documents(chunks, self.embeddings)
        except Exception as e:
            st.error(f"Vektör veritabanı oluşturulurken hata: {e}")
            self.vectordb = None

    def answer_question(self, question: str, model_name: str, temperature: float) -> dict:
        """Soruyu seçilen yerel modelle yanıtlar ve süreyi (hızı) ölçer."""
        if not self.documents:
            return {"answer": "❌ Lütfen önce doküman yükleyin.", "time": 0}
        if not self.vectordb:
            return {"answer": "⚠️ Arama dizini hazır değil.", "time": 0}

        try:
            start_time = time.time()

            # Ollama ile yerel modeli çağırıyoruz
            llm = Ollama(model=model_name, temperature=temperature)

            # Vektör veritabanından ilgili kısımları getir (Kaynak gösterme için)
            docs = self.vectordb.similarity_search(question, k=SIMILARITY_TOP_K)
            if not docs:
                return {"answer": "🔍 Dokümanlarda ilgili bir içerik bulunamadı.", "time": 0}

            context = "\n\n".join(
                [f"📄 Kaynak: {d.metadata.get('source', 'Bilinmeyen Doküman')}\n{d.page_content}" for d in docs]
            )

            # Zinciri çalıştır
            chain = self.prompt_template | llm
            response = chain.invoke({"context": context, "question": question})

            elapsed_time = round(time.time() - start_time, 2)

            # Kaynak dokümanların isimlerini listele
            sources = list(set([d.metadata.get('source', 'Bilinmeyen') for d in docs]))
            source_text = "\n\n**Dayanılan Kaynaklar:** " + ", ".join(sources)

            return {"answer": response + source_text, "time": elapsed_time}
        except Exception as e:
            return {
                "answer": f"❌ Yerel model çalıştırılırken hata oluştu. Ollama'nın açık ve modelin kurulu olduğundan emin olun. Hata: {str(e)}",
                "time": 0}

    def get_stats(self):
        total_files = len(self.processed_files)
        total_words = sum(info.get("word_count", 0) for info in self.processed_files.values())
        total_size = sum(info.get("size", 0) for info in self.processed_files.values())
        return {"files": total_files, "words": total_words, "size_mb": round(total_size / (1024 * 1024), 2)}


# ----------------- METNİN ÇIKARILMASI (OCR DAHİL) ----------------- #
def get_file_hash(file_content: bytes) -> str:
    return hashlib.md5(file_content).hexdigest()


def normalize_text(text: str) -> str:
    return " ".join(text.split())


def ocr_image(image_bytes: bytes) -> str:
    try:
        image = Image.open(io.BytesIO(image_bytes))
        return pytesseract.image_to_string(image, lang="tur+eng")
    except Exception:
        return ""


def extract_text_from_pdf(file_bytes: io.BytesIO, use_ocr: bool = True) -> str:
    import fitz
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        page_text = page.get_text()
        if page_text:
            text += page_text + "\n"
        if use_ocr:
            for img_meta in page.get_images(full=True):
                try:
                    xref = img_meta[0]
                    base_image = doc.extract_image(xref)
                    ocr_text = ocr_image(base_image["image"])
                    if ocr_text.strip() and ocr_text not in text:
                        text += "\n" + ocr_text
                except Exception:
                    continue
    return normalize_text(text)


def extract_text_from_docx(file_bytes: io.BytesIO, use_ocr: bool = True) -> str:
    from docx import Document as DocxDocument
    doc = DocxDocument(file_bytes)
    text = "\n".join(para.text for para in doc.paragraphs)
    if use_ocr:
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    ocr_text = ocr_image(rel.target_part.blob)
                    if ocr_text.strip() and ocr_text not in text:
                        text += "\n" + ocr_text
        except Exception:
            pass
    return normalize_text(text)


def extract_text_from_txt(file_bytes: io.BytesIO) -> str:
    raw = file_bytes.read()
    try:
        return normalize_text(raw.decode("utf-8"))
    except UnicodeDecodeError:
        return normalize_text(raw.decode("utf-8", errors="ignore"))


# ----------------- SESSION STATE BAŞLATMA ----------------- #
if "doc_manager" not in st.session_state:
    st.session_state.doc_manager = DocumentManager()
if "messages" not in st.session_state:
    st.session_state.messages = []

# ----------------- ARAYÜZ TASARIMI ----------------- #
st.markdown(
    """
<div class="main-header">
    <h1>🤖 Offline Kullanılabilen Yapay Zeka Doküman Asistanı</h1>
    <p style="font-size: 1.1em; opacity: 0.9;">
        Seçenek 4: Gizlilik Odaklı, Cihaz Üzerinde Çalışan Model Karşılaştırma Sistemi
    </p>
</div>
""",
    unsafe_allow_html=True,
)

# YAN MENÜ (SIDEBAR): Model Seçimi, Yapılandırma ve Dosya Yükleme
with st.sidebar:
    st.markdown("### 🤖 1. Yerel Model Yapılandırması")

    # YÖNERGE ŞARTI: En az 2 modelin karşılaştırılması
    selected_model = st.selectbox(
        "Kullanılacak Yapay Zekâ Modeli:",
        ["qwen2.5:1.5b", "phi3"],
        help="2 farklı yerel model kullanılabilir. (Hız ve doğruluk açısından farklılık gösterebilir.)"
    )

    # YÖNERGE ŞARTI: Prompt tasarımı ve model parametreleri kontrolü
    model_temperature = st.slider(
        "Model Yaratıcılığı (Temperature):",
        min_value=0.0, max_value=1.0, value=0.1, step=0.1,
        help="Düşük değerler dökümana daha sadık ve tutarlı cevaplar üretir."
    )

    st.markdown("---")
    st.markdown("### 📁 2. Doküman Yükleme")
    uploaded_files = st.file_uploader(
        "Dosyalarınızı seçin", type=["pdf", "docx", "txt"], accept_multiple_files=True
    )
    use_ocr = st.checkbox("🔍 Görseller için OCR Etkinleştir", value=True)

    if st.button("🗑️ Tüm Belleği Temizle"):
        st.session_state.doc_manager = DocumentManager()
        st.session_state.messages = []
        st.success("Sistem sıfırlandı!")
        st.rerun()

    # İstatistik Paneli
    stats = st.session_state.doc_manager.get_stats()
    if stats["files"] > 0:
        st.markdown("---")
        st.markdown("### 📊 Yüklenen Doküman İstatistikleri")
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 Dosya", stats["files"])
            st.metric("💾 Boyut (MB)", stats["size_mb"])
        with col2:
            st.metric("📝 Kelime", f"{stats['words']:,}")

# Yüklenen dosyaları işleme mantığı
if uploaded_files:
    progress_bar = st.progress(0)
    status_text = st.empty()
    for i, uploaded_file in enumerate(uploaded_files):
        progress_bar.progress((i + 1) / len(uploaded_files))
        status_text.text(f"İşleniyor: {uploaded_file.name}...")

        file_data = uploaded_file.getvalue()
        file_bytes = io.BytesIO(file_data)
        file_hash = get_file_hash(file_data)

        if file_hash in st.session_state.doc_manager.processed_files:
            continue

        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        text_content = ""
        try:
            if file_extension == ".pdf":
                text_content = extract_text_from_pdf(file_bytes, use_ocr)
            elif file_extension == ".docx":
                text_content = extract_text_from_docx(file_bytes, use_ocr)
            elif file_extension == ".txt":
                text_content = extract_text_from_txt(file_bytes)

            success, message = st.session_state.doc_manager.add_file(
                uploaded_file.name, text_content, file_hash, uploaded_file.size
            )
            if success:
                st.sidebar.success(message)
        except Exception as e:
            st.sidebar.error(f"❌ Hata {uploaded_file.name}: {str(e)}")

    with st.spinner("🔄 Yerel arama dizini (Vektör Veritabanı) güncelleniyor..."):
        st.session_state.doc_manager._rebuild_vectordb()
    progress_bar.empty()
    status_text.empty()

# ANA SOHBET EKRANI
if st.session_state.doc_manager.documents:
    st.markdown(f"### 💬 Dokümanlarınız ile Mesajlaşın ({selected_model} aktif)")

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    if prompt := st.chat_input("Dokümanlarınız hakkında soru sorun..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.chat_message("assistant"):
            with st.spinner("🤔 Yerel model yanıt üretiyor..."):
                result = st.session_state.doc_manager.answer_question(
                    prompt, model_name=selected_model, temperature=model_temperature
                )
            st.markdown(result["answer"])
            # YÖNERGE ŞARTI: Hız/performans metriğinin kullanıcıya gösterilmesi
            st.caption(
                f"⏱️ Yanıt Süresi: {result['time']} saniye | Model: {selected_model} | Sıcaklık: {model_temperature}")

        st.session_state.messages.append({"role": "assistant", "content": result["answer"]})
else:
    st.markdown(
        """
    <div class="feature-card">
        <h3>🚀 Sistem Nasıl Çalışır?</h3>
        <p>Uygulama, Seçenek 4 asgari teknik gereksinimlerine göre tamamen yerel cihaz mimarisine uygun olarak tasarlanmıştır:</p>
        <ol>
            <li><strong>Veri Gizliliği:</strong> Yüklediğiniz dokümanlar hiçbir bulut servisine veya dış API'ye gönderilmez.</li>
            <li><strong>Model Karşılaştırması:</strong> Sol menüden <code>qwen2.5:1.5b</code> veya <code>phi3</code> modellerini seçerek sistemin hız/doğruluk performansını analiz edebilirsiniz.</li>
            <li><strong>Parametre Özelleştirme:</strong> Temperature ayarı ile yerel modelin prompt çıktı dinamiklerini değiştirebilirsiniz.</li>
        </ol>
    </div>
    """, unsafe_allow_html=True)


    